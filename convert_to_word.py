from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 标题
title = doc.add_heading('掌中光项目 Git 协作指南（新手版）', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('仓库地址：https://github.com/shuzhifei/zhangzhongguang')
doc.add_paragraph('创作者 GitHub：shuzhifei')
doc.add_paragraph('项目名：掌中光 - 皮影戏AI互动影游')
doc.add_paragraph('')

# ---- 一 ----
doc.add_heading('一、这个页面是什么？', level=1)
p = doc.add_paragraph()
p.add_run('这是 GitHub 仓库的主页面，相当于你的代码"网盘主页"。').font.size = Pt(11)
doc.add_paragraph('页面名称：zhangzhongguang（公开）')
doc.add_paragraph('网址：https://github.com/shuzhifei/zhangzhongguang')
doc.add_paragraph('Public：公开仓库，所有人可见')
doc.add_paragraph('Code：显示代码文件')
doc.add_paragraph('Add collaborators：添加协作者（把队友加进来才能一起编辑）')

# ---- 二 ----
doc.add_heading('二、第一步：把自己加进仓库（创作者操作）', level=1)
doc.add_paragraph('只有加了仓库权限，队友才能上传代码。创作者需要做：')
doc.add_paragraph('在页面右上角找到"Add collaborators"')
doc.add_paragraph('点击"Invite collaborators"')
doc.add_paragraph('输入队友的 GitHub 用户名或邮箱')
doc.add_paragraph('发送邀请')
doc.add_paragraph('队友收到邮件后点击接受，就有了编辑权限。')

# ---- 三 ----
doc.add_heading('三、第二步：队友下载代码（每人只执行一次）', level=1)

doc.add_heading('1. 安装 Git', level=2)
doc.add_paragraph('访问 https://git-scm.com/download/win')
doc.add_paragraph('下载安装包，一路点击"下一步"安装。')

doc.add_heading('2. 设置身份（只用设一次）', level=2)
doc.add_paragraph('打开 VS Code，按 Ctrl + ` （ESC 下面那个键）打开终端，输入：')
code1 = doc.add_paragraph()
run = code1.add_run('git config --global user.name "你的姓名"\ngit config --global user.email "你的QQ号@qq.com"')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('解释：告诉 Git 你是谁，以后提交代码时会显示你的名字。')

doc.add_heading('3. 下载代码', level=2)
doc.add_paragraph('在终端中输入：')
code2 = doc.add_paragraph()
run = code2.add_run('cd D:\\\ngit clone git@github.com:shuzhifei/zhangzhongguang.git')
run.font.name = 'Consolas'
run.font.size = Pt(10)

p = doc.add_paragraph()
p.add_run('cd D:\\ ').bold = True
p.add_run('—— 进入 D 盘')
p = doc.add_paragraph()
p.add_run('git clone').bold = True
p.add_run(' —— 从 GitHub 下载整个项目到本地')
p = doc.add_paragraph()
p.add_run('git@github.com:...').bold = True
p.add_run(' —— 用 SSH 地址下载（国内网络更快）')

doc.add_paragraph('执行完后，队友的电脑上会出现一个 D:\\zhangzhongguang 文件夹，里面就有全部代码。')

# ---- 四 ----
doc.add_heading('四、第三步：修改代码后上传（日常操作）', level=1)

doc.add_heading('1. 查看改动（可选）', level=2)
code3 = doc.add_paragraph()
run = code3.add_run('cd D:\\zhangzhongguang\ngit status')
run.font.name = 'Consolas'
run.font.size = Pt(10)
doc.add_paragraph('解释：查看哪些文件被修改了，显示红色表示还没提交。')

doc.add_heading('2. 提交改动', level=2)
p = doc.add_paragraph()
p.add_run('git add .').bold = True
p = doc.add_paragraph()
p.add_run('解释：把修改过的文件放进"暂存区"（就像把东西放进购物车）。')

p = doc.add_paragraph()
p.add_run('git commit -m "你的改动说明"').bold = True
p = doc.add_paragraph()
p.add_run('解释：正式存档，-m 后面写你做了什么（比如"A同学：修复了皮影拖拽的Bug"）。')

doc.add_heading('3. 上传到 GitHub', level=2)
p = doc.add_paragraph()
p.add_run('git push').bold = True
p = doc.add_paragraph()
p.add_run('解释：把本地存档推送到 GitHub 上，让其他人能看到你的代码。')

# ---- 五 ----
doc.add_heading('五、创作者拉取队友的代码', level=1)
p = doc.add_paragraph()
p.add_run('git pull').bold = True
p = doc.add_paragraph()
p.add_run('解释：把 GitHub 上队友写的新代码拉取到本地电脑。')

# ---- 六 ----
doc.add_heading('六、完整协作流程', level=1)
doc.add_paragraph('创作者改代码 → git add → git commit → git push')
doc.add_paragraph('   ↓')
doc.add_paragraph('GitHub 仓库 ← 更新代码')
doc.add_paragraph('   ↓')
doc.add_paragraph('队友：git clone（第一次）→ git pull（后续）→ 修改代码 → git push')
doc.add_paragraph('   ↓')
doc.add_paragraph('创作者：git pull → 收到队友代码')

# ---- 七 ----
doc.add_heading('七、常见问题', level=1)

doc.add_heading('Q1：输入 git clone 后报错"Permission denied"', level=2)
doc.add_paragraph('说明没有被邀请加入仓库。请创作者在 GitHub 仓库页面点"Add collaborators"，输入你的 GitHub 用户名，邀请你。')

doc.add_heading('Q2：git push 报错"failed to push"', level=2)
doc.add_paragraph('可能是你的本地代码不是最新的。先执行 git pull，再 git push。')

doc.add_heading('Q3：两个人同时改了同一个文件怎么办？', level=2)
doc.add_paragraph('Git 会提示冲突。打开冲突文件，手动选择保留哪部分代码，然后重新提交。')

doc.add_heading('Q4：怎么查看仓库里有哪些文件？', level=2)
doc.add_paragraph('在浏览器打开 https://github.com/shuzhifei/zhangzhongguang，点击文件夹即可查看。')

# ---- 八 ----
doc.add_heading('八、重要命令速查表', level=1)

table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
headers = ['命令', '作用', '使用频率']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['git clone', '下载仓库', '每人一次'],
    ['git pull', '拉取最新代码', '每次开始工作前'],
    ['git add .', '把改动放进购物车', '每次提交前'],
    ['git commit -m "说明"', '存档（提交到本地）', '每次提交前'],
    ['git push', '上传到 GitHub', '每次提交后'],
    ['git status', '查看当前状态', '随时'],
    ['git log', '查看提交历史', '偶尔'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = cell_text

doc.add_paragraph('')
doc.add_paragraph('仓库地址：https://github.com/shuzhifei/zhangzhongguang')
doc.add_paragraph('创作者 GitHub：shuzhifei')
doc.add_paragraph('项目名：掌中光 - 皮影戏AI互动影游')
doc.add_paragraph('祝协作愉快！')

doc.save(r'D:\掌中光\Git协作指南（新手版）.docx')
print("Word 文档已生成：D:\\掌中光\\Git协作指南（新手版）.docx")
